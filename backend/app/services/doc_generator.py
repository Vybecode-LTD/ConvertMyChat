"""DOCX export generator — renders Markdown content as a styled Word document."""

import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import ConversationData

_INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`", re.DOTALL)


def _apply_inline(para, text: str):
    """Add runs to a paragraph with **bold**, *italic*, and `code` support."""
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1) is not None:
            para.add_run(m.group(1)).bold = True
        elif m.group(2) is not None:
            para.add_run(m.group(2)).italic = True
        elif m.group(3) is not None:
            r = para.add_run(m.group(3))
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx]
        for c_idx in range(ncols):
            cell = tr.cells[c_idx]
            txt = row_data[c_idx] if c_idx < len(row_data) else ""
            run = cell.paragraphs[0].add_run(txt)
            if r_idx == 0:
                run.bold = True
    doc.add_paragraph()


def _render_markdown(doc: Document, text: str):
    """Render a markdown string into the document."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph("\n".join(code_lines))
                para.paragraph_format.left_indent = Cm(0.5)
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # Markdown table
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl.strip()):
                    continue
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                if cells:
                    rows.append(cells)
            _add_table(doc, rows)
            continue

        # ATX heading
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 6))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-*+]\s+(.+)$", line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            _apply_inline(para, m.group(1))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.+)$", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            _apply_inline(para, m.group(1))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _apply_inline(para, line)
        i += 1


def generate_docx(conversation: ConversationData) -> bytes:
    doc = Document()

    doc.add_heading(conversation.title, level=1)

    meta = doc.add_paragraph()
    r = meta.add_run(
        f"Source: {conversation.share_url}\n"
        f"Extracted: {conversation.extracted_at.strftime('%Y-%m-%d %H:%M UTC')} · "
        f"{conversation.message_count} messages"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.font.italic = True

    doc.add_paragraph()

    for msg in conversation.messages:
        label = "User" if msg.role == "user" else "Gemini"
        doc.add_heading(label, level=3)
        _render_markdown(doc, msg.content)
        doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Exported by ConvertMyChat")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
